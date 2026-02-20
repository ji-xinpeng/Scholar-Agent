import os
import uuid
from docx import Document
import PyPDF2
from datetime import datetime, timezone
from typing import List, Optional
from app.core.config import settings
from app.infrastructure.storage.database.connection import get_db
from app.infrastructure.logging.config import logger


class DocumentService:
    def __init__(self):
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    def _get_user_upload_dir(self, user_id: str) -> str:
        """获取用户上传目录"""
        user_dir = os.path.join(settings.UPLOAD_DIR, user_id)
        os.makedirs(user_dir, exist_ok=True)
        return user_dir

    def create_document(self, user_id: str, filename: str, content: str = "", file_type: str = "markdown", folder_id: Optional[str] = None) -> dict:
        """创建新文档（支持 markdown、text、word 格式）"""
        db = get_db()
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()

        ext_map = {"markdown": ".md", "text": ".txt", "word": ".docx"}
        ext = ext_map.get(file_type, ".md")
        stored_name = f"{doc_id}{ext}"
        user_dir = self._get_user_upload_dir(user_id)
        file_path = os.path.join(user_dir, stored_name)

        try:
            if file_type == "word":
                docx = Document()
                for line in content.split("\n"):
                    docx.add_paragraph(line)
                docx.save(file_path)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)

            file_size = os.path.getsize(file_path)
            db.execute(
                "INSERT INTO documents (id, user_id, folder_id, filename, original_name, file_path, file_size, file_type, status, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (doc_id, user_id, folder_id, stored_name, filename, file_path, file_size, file_type, "uploaded", now)
            )
            db.commit()
            return {
                "id": doc_id, "user_id": user_id, "folder_id": folder_id,
                "filename": stored_name, "original_name": filename,
                "file_size": file_size, "file_type": file_type,
                "page_count": 0, "status": "uploaded", "created_at": now
            }
        except Exception as e:
            logger.error(f"创建文档失败: {e}")
            raise

    def upload_document(self, user_id: str, filename: str, file_content: bytes, folder_id: Optional[str] = None) -> dict:
        db = get_db()
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()

        ext = os.path.splitext(filename)[1].lower()
        file_type_map = {".pdf": "pdf", ".doc": "word", ".docx": "word", ".md": "markdown", ".txt": "text"}
        file_type = file_type_map.get(ext, "other")

        stored_name = f"{doc_id}{ext}"
        user_dir = self._get_user_upload_dir(user_id)
        file_path = os.path.join(user_dir, stored_name)
        with open(file_path, "wb") as f:
            f.write(file_content)

        file_size = len(file_content)
        db.execute(
            "INSERT INTO documents (id, user_id, folder_id, filename, original_name, file_path, file_size, file_type, status, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (doc_id, user_id, folder_id, stored_name, filename, file_path, file_size, file_type, "uploaded", now)
        )
        db.commit()
        return {
            "id": doc_id, "user_id": user_id, "folder_id": folder_id,
            "filename": stored_name, "original_name": filename,
            "file_size": file_size, "file_type": file_type,
            "page_count": 0, "status": "uploaded", "created_at": now
        }

    def list_documents(self, user_id: str, folder_id: Optional[str] = None, page: int = 1, page_size: int = 10) -> dict:
        db = get_db()
        offset = (page - 1) * page_size

        if folder_id:
            count = db.execute("SELECT COUNT(*) FROM documents WHERE user_id = ? AND folder_id = ?", (user_id, folder_id)).fetchone()[0]
            rows = db.execute(
                "SELECT * FROM documents WHERE user_id = ? AND folder_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?",
                (user_id, folder_id, page_size, offset)
            ).fetchall()
        else:
            count = db.execute("SELECT COUNT(*) FROM documents WHERE user_id = ?", (user_id,)).fetchone()[0]
            rows = db.execute(
                "SELECT * FROM documents WHERE user_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?",
                (user_id, page_size, offset)
            ).fetchall()

        documents = [dict(r) for r in rows]
        return {"documents": documents, "total": count, "page": page, "page_size": page_size}

    def get_document(self, doc_id: str) -> Optional[dict]:
        db = get_db()
        row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(row) if row else None
    
    def find_existing_paper(self, user_id: str, pdf_url: Optional[str] = None, title: Optional[str] = None) -> Optional[dict]:
        """查找是否已存在相同的论文（通过原始文件名匹配）"""
        db = get_db()
        
        if not title:
            return None
        
        safe_title = "".join([c if c.isalnum() or c in (' ', '-', '_') else '_' for c in title])
        filename_pattern = f"%{safe_title[:30]}%.pdf"
        
        rows = db.execute(
            "SELECT * FROM documents WHERE user_id = ? AND original_name LIKE ? ORDER BY created_at DESC LIMIT 5",
            (user_id, filename_pattern)
        ).fetchall()
        
        if rows:
            for row in rows:
                doc = dict(row)
                original_name = doc.get("original_name", "")
                if safe_title.lower() in original_name.lower():
                    return doc
        
        return None

    def delete_document(self, doc_id: str) -> bool:
        db = get_db()
        doc = db.execute("SELECT file_path FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return False
        file_path = doc["file_path"]
        if os.path.exists(file_path):
            os.remove(file_path)
        db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        db.commit()
        return True

    def move_document(self, doc_id: str, folder_id: Optional[str]) -> bool:
        db = get_db()
        db.execute("UPDATE documents SET folder_id = ? WHERE id = ?", (folder_id, doc_id))
        db.commit()
        return True

    # ======== 文件夹 ========
    def create_folder(self, user_id: str, name: str, parent_id: Optional[str] = None) -> dict:
        db = get_db()
        folder_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "INSERT INTO folders (id, user_id, name, parent_id, created_at) VALUES (?, ?, ?, ?, ?)",
            (folder_id, user_id, name, parent_id, now)
        )
        db.commit()
        return {"id": folder_id, "user_id": user_id, "name": name, "parent_id": parent_id, "created_at": now, "document_count": 0}

    def list_folders(self, user_id: str) -> List[dict]:
        db = get_db()
        rows = db.execute("SELECT * FROM folders WHERE user_id = ? ORDER BY created_at", (user_id,)).fetchall()
        folders = []
        for r in rows:
            f = dict(r)
            f["document_count"] = db.execute(
                "SELECT COUNT(*) FROM documents WHERE folder_id = ?", (f["id"],)
            ).fetchone()[0]
            folders.append(f)
        return folders

    def delete_folder(self, folder_id: str) -> bool:
        db = get_db()
        db.execute("UPDATE documents SET folder_id = NULL WHERE folder_id = ?", (folder_id,))
        db.execute("DELETE FROM folders WHERE id = ?", (folder_id,))
        db.commit()
        return True


    def get_document_content(self, doc_id: str) -> Optional[str]:
        doc = self.get_document(doc_id)
        if not doc:
            return None
        
        file_path = doc["file_path"]
        if not os.path.exists(file_path):
            return None
        
        file_type = doc["file_type"]
        content = ""
        
        try:
            if file_type == "pdf":
                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
            elif file_type == "word":
                docx = Document(file_path)
                for para in docx.paragraphs:
                    content += para.text + "\n"
            elif file_type in ["markdown", "text"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error reading document {doc_id}: {e}")
            return None
        
        return content

    def append_document_content(self, doc_id: str, content_to_append: str) -> bool:
        """在文档末尾追加内容"""
        existing = self.get_document_content(doc_id)
        if existing is None:
            return False
        new_content = existing.rstrip("\n") + "\n" + content_to_append
        return self.update_document_content(doc_id, new_content)

    def replace_document_content(self, doc_id: str, old_text: str, new_text: str) -> bool:
        """
        替换文档中的指定文本
        支持多种匹配方式，提高成功率
        """
        existing = self.get_document_content(doc_id)
        if existing is None:
            logger.error(f"文档不存在: {doc_id}")
            return False

        logger.debug(f"尝试替换文档内容")
        logger.debug(f"old_text 长度: {len(old_text)}")
        logger.debug(f"old_text 前200字符: {old_text[:200]}")
        logger.debug(f"文档内容长度: {len(existing)}")

        # 1. 精确匹配
        if old_text in existing:
            new_content = existing.replace(old_text, new_text, 1)
            logger.info(f"精确匹配成功，正在更新文档")
            return self.update_document_content(doc_id, new_content)

        # 2. 尝试去除首尾空格匹配
        old_text_stripped = old_text.strip()
        if old_text_stripped and old_text_stripped in existing:
            # 查找包含这个 stripped 文本的完整段落
            lines = existing.split("\n")
            found_idx = -1
            for i, line in enumerate(lines):
                if old_text_stripped in line:
                    found_idx = i
                    break
            
            if found_idx != -1:
                # 尝试找到前后更多上下文进行匹配
                start = max(0, found_idx - 5)
                end = min(len(lines), found_idx + 6)
                candidate = "\n".join(lines[start:end])
                if old_text_stripped in candidate:
                    logger.info(f"通过 stripped 文本 + 上下文匹配成功")
                    new_content = existing.replace(candidate, new_text, 1)
                    return self.update_document_content(doc_id, new_content)

        # 3. 如果都匹配失败，记录详细日志
        logger.warning(f"无法匹配要替换的文本")
        logger.warning(f"old_text: {repr(old_text)}")
        logger.warning(f"old_text_stripped: {repr(old_text_stripped)}")
        
        return False

    def update_document_content(self, doc_id: str, content: str) -> bool:
        """更新文档内容（支持 word、markdown、text 格式）"""
        doc = self.get_document(doc_id)
        if not doc:
            return False

        file_path = doc["file_path"]
        file_type = doc["file_type"]

        if not os.path.exists(file_path):
            return False

        try:
            if file_type == "word":
                docx = Document()
                for line in content.split("\n"):
                    docx.add_paragraph(line)
                docx.save(file_path)
            elif file_type in ["markdown", "text"]:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)
            else:
                logger.warning(f"不支持编辑的文件类型: {file_type}")
                return False

            return True
        except Exception as e:
            logger.error(f"更新文档 {doc_id} 失败: {e}")
            return False

    async def parse_document(self, doc_id: str) -> bool:

        return True


document_service = DocumentService()
