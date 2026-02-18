import os
import uuid
import shutil
from datetime import datetime, timezone
from typing import List, Optional
from app.core.config import settings
from app.core.database import get_db
from app.core.logger import logger

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger.warning("PyPDF2 not installed, PDF parsing disabled")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed, DOCX parsing disabled")


class DocumentService:
    def __init__(self):
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    def upload_document(self, user_id: str, filename: str, file_content: bytes, folder_id: Optional[str] = None) -> dict:
        db = get_db()
        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()

        ext = os.path.splitext(filename)[1].lower()
        file_type_map = {".pdf": "pdf", ".doc": "word", ".docx": "word", ".md": "markdown", ".txt": "text"}
        file_type = file_type_map.get(ext, "other")

        stored_name = f"{doc_id}{ext}"
        file_path = os.path.join(settings.UPLOAD_DIR, stored_name)
        with open(file_path, "wb") as f:
            f.write(file_content)

        file_size = len(file_content)
        db.execute(
            "INSERT INTO documents (id, user_id, folder_id, filename, original_name, file_path, file_size, file_type, status, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (doc_id, user_id, folder_id, stored_name, filename, file_path, file_size, file_type, "uploaded", now)
        )
        db.commit()
        return {
            "id": doc_id, "user_id": user_id, "folder_id": folder_id,
            "filename": stored_name, "original_name": filename,
            "file_size": file_size, "file_type": file_type,
            "page_count": 0, "status": "uploaded", "created_at": now
        }

    def list_documents(self, user_id: str, folder_id: Optional[str] = None, page: int = 1, page_size: int = 10) -> dict:
        db = get_db()
        offset = (page - 1) * page_size

        if folder_id:
            count = db.execute("SELECT COUNT(*) FROM documents WHERE user_id = ? AND folder_id = ?", (user_id, folder_id)).fetchone()[0]
            rows = db.execute(
                "SELECT * FROM documents WHERE user_id = ? AND folder_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?",
                (user_id, folder_id, page_size, offset)
            ).fetchall()
        else:
            count = db.execute("SELECT COUNT(*) FROM documents WHERE user_id = ?", (user_id,)).fetchone()[0]
            rows = db.execute(
                "SELECT * FROM documents WHERE user_id = ? ORDER BY created_at DESC LIMIT ? OFFSET ?",
                (user_id, page_size, offset)
            ).fetchall()

        documents = [dict(r) for r in rows]
        return {"documents": documents, "total": count, "page": page, "page_size": page_size}

    def get_document(self, doc_id: str) -> Optional[dict]:
        db = get_db()
        row = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(row) if row else None

    def delete_document(self, doc_id: str) -> bool:
        db = get_db()
        doc = db.execute("SELECT file_path FROM documents WHERE id = ?", (doc_id,)).fetchone()
        if not doc:
            return False
        file_path = doc["file_path"]
        if os.path.exists(file_path):
            os.remove(file_path)
        db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        db.commit()
        return True

    def move_document(self, doc_id: str, folder_id: Optional[str]) -> bool:
        db = get_db()
        db.execute("UPDATE documents SET folder_id = ? WHERE id = ?", (folder_id, doc_id))
        db.commit()
        return True

    # ======== 文件夹 ========
    def create_folder(self, user_id: str, name: str, parent_id: Optional[str] = None) -> dict:
        db = get_db()
        folder_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        db.execute(
            "INSERT INTO folders (id, user_id, name, parent_id, created_at) VALUES (?, ?, ?, ?, ?)",
            (folder_id, user_id, name, parent_id, now)
        )
        db.commit()
        return {"id": folder_id, "user_id": user_id, "name": name, "parent_id": parent_id, "created_at": now, "document_count": 0}

    def list_folders(self, user_id: str) -> List[dict]:
        db = get_db()
        rows = db.execute("SELECT * FROM folders WHERE user_id = ? ORDER BY created_at", (user_id,)).fetchall()
        folders = []
        for r in rows:
            f = dict(r)
            f["document_count"] = db.execute(
                "SELECT COUNT(*) FROM documents WHERE folder_id = ?", (f["id"],)
            ).fetchone()[0]
            folders.append(f)
        return folders

    def delete_folder(self, folder_id: str) -> bool:
        db = get_db()
        db.execute("UPDATE documents SET folder_id = NULL WHERE folder_id = ?", (folder_id,))
        db.execute("DELETE FROM folders WHERE id = ?", (folder_id,))
        db.commit()
        return True


    def get_document_content(self, doc_id: str) -> Optional[str]:
        doc = self.get_document(doc_id)
        if not doc:
            return None
        
        file_path = doc["file_path"]
        if not os.path.exists(file_path):
            return None
        
        file_type = doc["file_type"]
        content = ""
        
        try:
            if file_type == "pdf" and HAS_PDF:
                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
            elif file_type == "word" and HAS_DOCX:
                docx = Document(file_path)
                for para in docx.paragraphs:
                    content += para.text + "\n"
            elif file_type in ["markdown", "text"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error reading document {doc_id}: {e}")
            return None
        
        return content

    def append_document_content(self, doc_id: str, content_to_append: str) -> bool:
        """在文档末尾追加内容"""
        existing = self.get_document_content(doc_id)
        if existing is None:
            return False
        new_content = existing.rstrip("\n") + "\n" + content_to_append
        return self.update_document_content(doc_id, new_content)

    def replace_document_content(self, doc_id: str, old_text: str, new_text: str) -> bool:
        """替换文档中的指定文本"""
        existing = self.get_document_content(doc_id)
        if existing is None:
            return False
        if old_text not in existing:
            return False
        new_content = existing.replace(old_text, new_text, 1)
        return self.update_document_content(doc_id, new_content)

    def update_document_content(self, doc_id: str, content: str) -> bool:
        """更新文档内容（支持 word、markdown、text 格式）"""
        doc = self.get_document(doc_id)
        if not doc:
            return False

        file_path = doc["file_path"]
        file_type = doc["file_type"]

        if not os.path.exists(file_path):
            return False

        try:
            if file_type == "word" and HAS_DOCX:
                docx = Document()
                for line in content.split("\n"):
                    docx.add_paragraph(line)
                docx.save(file_path)
            elif file_type in ["markdown", "text"]:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)
            else:
                logger.warning(f"不支持编辑的文件类型: {file_type}")
                return False

            return True
        except Exception as e:
            logger.error(f"更新文档 {doc_id} 失败: {e}")
            return False

    async def parse_document(self, doc_id: str) -> bool:

        return True


document_service = DocumentService()
